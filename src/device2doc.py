#############################
#
#    copyright 2016 Open Interconnect Consortium, Inc. All rights reserved.
#    Redistribution and use in source and binary forms, with or without modification,
#    are permitted provided that the following conditions are met:
#    1.  Redistributions of source code must retain the above copyright notice,
#        this list of conditions and the following disclaimer.
#    2.  Redistributions in binary form must reproduce the above copyright notice,
#        this list of conditions and the following disclaimer in the documentation and/or other materials provided
#        with the distribution.
#         
#    THIS SOFTWARE IS PROVIDED BY THE OPEN INTERCONNECT CONSORTIUM, INC. "AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES,
#    INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE OR
#    WARRANTIES OF NON-INFRINGEMENT, ARE DISCLAIMED. IN NO EVENT SHALL THE OPEN INTERCONNECT CONSORTIUM, INC. OR
#    CONTRIBUTORS BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
#    (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS;
#    OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
#    OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE,
#    EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
#############################


import time 
import os    
import json
import random
import sys
import argparse
import traceback
from datetime import datetime
from time import gmtime, strftime
import jsonref

if sys.version_info < (3, 5):
    raise Exception("ERROR: Python 3.5 or more is required, you are currently running Python %d.%d!" %
                    (sys.version_info[0], sys.version_info[1]))
#
# docx imports
#
try: 
    from docx import Document
except:
    print("missing docx:")
    print ("Trying to Install required module: python-docx (docx)")
    os.system('python3 -m pip install python-docx')
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_json_schema(filename, my_dir):
    """
    load the JSON schema file
    :param filename: filename (with extension)
    :param my_dir: path to the file
    :return: json_dict
    """
    full_path = os.path.join(my_dir, filename)
    if os.path.isfile(full_path) is False:
        print ("json file does not exist:", full_path)
            
    linestring = open(full_path, 'r').read()
    json_dict = json.loads(linestring)

    return json_dict


def get_dir_list(dir, ext=None):
    """
    get all files (none recursive) in the specified dir
    :param dir: path to the directory
    :param ext: filter on extension
    :return: list of files (only base_name)
    """
    only_files = [f for f in listdir(dir) if isfile(join(dir, f))]
    # remove .bak files
    new_list = [x for x in only_files if not x.endswith(".bak")]
    if ext is not None:
        cur_list = new_list
        new_list = [x for x in cur_list if x.endswith(ext)]
    return new_list
    
    
def find_key(rec_dict, target, depth=0):
    """
    find key "target" in recursive dict
    :param rec_dict: dict to search in, json schema dict, so it is combination of dict and arrays
    :param target: target key to search for
    :param depth: depth of the search (recursion)
    :return:
    """
    try:
        if isinstance(rec_dict, dict):
            for key, value in rec_dict.items():
                if key == target:
                    return rec_dict[key]
            for key, value in rec_dict.items():
                r = find_key(value, target, depth+1)
                if r is not None:
                        return r
        #else:
        #    print ("no dict:", rec_dict)
    except:
        traceback.print_exc()


def find_key_link(rec_dict, target, depth=0):
    """
    find the first key recursively
    also traverse lists (arrays, oneOf,..) but only returns the first occurance
    :param rec_dict: dict to search in, json schema dict, so it is combination of dict and arrays
    :param target: target key to search for
    :param depth: depth of the search (recursion)
    :return:
    """
    if isinstance(rec_dict, dict):
        # direct key
        for key, value in rec_dict.items():
            if key == target:
                return rec_dict[key]
        # key is in array
        rvalues = []
        found = False
        for key, value in rec_dict.items():
            if key in ["oneOf", "allOf", "anyOf"]:
                for val in value:
                    if val == target:
                        return val
                    if isinstance(val, dict):
                        r = find_key_link(val, target, depth+1)
                        if r is not None:
                            found = True
                            # TODO: this should return an array, now it only returns the last found item
                            rvalues = r
        if found:
            return rvalues
        # key is an dict
        for key, value in rec_dict.items():
            r = find_key_link(value, target, depth+1)
            if r is not None:
                return r #[list(r.items())]


class CreateWordDoc(object):
    def __init__(self, device=None, lbnldevice=None, docx_name_in=None, docx_name_out=None, resource_name=None):
        """
        initialize the class


        """
        # input arguments
        self.docx_name_in = docx_name_in
        self.docx_name_out = docx_name_out
        
        # initialise the variable
        self.device = device
        self.device_filename = device
        self.lbnldevice = lbnldevice
        if lbnldevice is not None:
            self.device_filename = lbnldevice

        schema_string = open(self.device_filename, 'r').read()   
        json_dict = json.loads(schema_string)
        self.json_parse_tree = json_dict
        
    def swag_sanitize_description(self, description):
        """
        removes line breaks, quotes  etc
        :param description: input string
        :return: text string
        """
        text = description.replace("\n", "@cr").replace("'", "<COMMA>").replace('"', "<COMMA>")
        return text
        
    def swag_unsanitize_description(self, description):
        """
        removes line breaks, quotes  etc
        :param description: input string
        :return: text string
        """
        text = description.replace("@cr", "\n").replace("<COMMA>", "'")
        return text
    
    def resources_per_device(self, parse_tree):
        """
        list all properties (attributes) in an table.
        create the table and fill it up
        :param parse_tree: json parse_tree of the device list 
        """
        self.tableAttribute = self.document.add_table(rows=1, cols=4, style='TABLE-A')
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = 'Device Name\n (informative)'
        hdr_cells[1].text = 'Device Type ("rt")\n (Normative)'
        hdr_cells[2].text = 'Required Resource Name'
        hdr_cells[3].text = 'Required Resoure Type'

        for device_data in parse_tree:
            row_cells = self.tableAttribute.add_row().cells
            row_cells[0].text = device_data["devicename"]
            row_cells[1].text = device_data["devicetype"]
            resources = device_data["resources"]
            total_resources = len(resources)
            counter = 1
            for resource in resources:
                if counter == 1:
                    for resource in resources:
                        row_cells[2].text = resource["resourcetypetitle"]
                        row_cells[3].text = resource["resourcetypeid"] 
                    counter += 1
                if len(resources) > 1:
                    row_cell_res = self.tableAttribute.add_row().cells
                    row_cell_res[2].text = resource["resourcetypetitle"]
                    row_cell_res[3].text = resource["resourcetypeid"] 
                    a, b = row_cell_res[:2]
                    a.merge(b)
                    
    def lbnl_resources_per_device(self, parse_tree):
        """
        list all properties (attributes) in an table.
        create the table and fill it up
        :param parse_tree: json parse_tree of the device list 
        """
        
        text = "Table was prepared at Lawrence Berkeley National Laboratory under Contract No. DEAC02-05CH11231 with the U.S. Department of Energy."
        self.document.add_paragraph(text)
        
        self.tableAttribute = self.document.add_table(rows=1, cols=6, style='TABLE-A')
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = 'Category and Device)'
        hdr_cells[1].text = 'Comment'
        # previous
        hdr_cells[2].text = 'Device Name\n (informative)'
        hdr_cells[3].text = 'Device Type ("rt")\n (Normative)'
        hdr_cells[4].text = 'Required Resource Name'
        hdr_cells[5].text = 'Required Resoure Type'
        
        
        for category in parse_tree:
            cat_cells = self.tableAttribute.add_row().cells
            print ("category:", category["category"] )
            #cat_cells[0].text = category["category"]
            cat_cells[0].paragraphs[0].add_run(category["category"]).bold = True
            # merge all... merge can only merge 2 cells into 1...
            a, b = cat_cells[:2]
            a.merge(b)
            #a, b = cat_cells[:2]
            #a.merge(b)
            #a, b = cat_cells[:2]
            #a.merge(b)
            #a, b = cat_cells[:2]
            #a.merge(b)
            
            for lnbldevice in category["devices"]:
                lnbl_cells = self.tableAttribute.add_row().cells
                lnbl_cells[0].text = str(lnbldevice["name"]) #+ "\n" + str(lnbldevice["rt"])
                lnbl_cells[1].text = str(lnbldevice["comment"])
                lnbl_cells[3].text = str(lnbldevice["rt"])
                for device_data in lnbldevice["exising"]:  
                    row_cells = self.tableAttribute.add_row().cells
                    a, b = row_cells[:2]
                    a.merge(b)
                    row_cells[2].text = device_data["devicename"]
                    row_cells[3].text = device_data["devicetype"]
                    resources = device_data["resources"]
                    total_resources = len(resources)
                    first = True
                    for resource in resources:
                        if first is True:
                            first = False
                            row_cells[4].text = resource["resourcetypetitle"]
                            row_cells[5].text = resource["resourcetypeid"] 
                        else:
                            row_cell_res = self.tableAttribute.add_row().cells
                            row_cell_res[4].text = resource["resourcetypetitle"]
                            row_cell_res[5].text = resource["resourcetypeid"] 
                            a, b = row_cell_res[2:4]
                            a.merge(b)
               
    
    def generate_sections(self, parse_tree):
        """
        generate the individual sections
        :param parse_tree:
        """
        self.list_devices(parse_tree)
    
    def convert(self):
        """
        conversion of the swagger data into the word document

        :return:
        """
        try:
            self.document = Document(docx=self.docx_name_in)
        except:
            print ("could not load file: ", self.docx_name_in)
            print ("make sure that docx file exist..")
            return

        if self.device is not None:
            self.resources_per_device(self.json_parse_tree)
        if self.lbnldevice is not None:
            self.lbnl_resources_per_device(self.json_parse_tree)
        ### add here more conversions going forward..
            
        if self.docx_name_out is not None:
            self.document.save(self.docx_name_out)
            print ("document saved..", self.docx_name_out)


#
#   main of script
#
print ("************************")
print ("*** device2doc (v1) ***")
print ("************************")
parser = argparse.ArgumentParser()

parser.add_argument( "-ver"        , "--verbose"    , help="Execute in verbose mode", action='store_true')

parser.add_argument( "-device"    , "--device"    , default=None,
                     help="device file name (json)",  nargs='?', const="", required=False)
parser.add_argument( "-lbnldevice"    , "--lbnldevice"    , default=None,
                     help="lbln device file name (json)",  nargs='?', const="", required=False)
parser.add_argument( "-docx"       , "--docx"       , default=None,
                     help="word file in",  nargs='?', const="", required=False)
parser.add_argument( "-word_out"   , "--word_out"   , default=None,
                     help="word file out",  nargs='?', const="", required=False)

args = parser.parse_args()


print("device file     : " + str(args.device))
print("lbnldevice file : " + str(args.lbnldevice))
print("docx            : " + str(args.docx))
print("word_out        : " + str(args.word_out))
print("")

try:
    if args.device is not None:
        worddoc = CreateWordDoc(device=args.device)
        worddoc.docx_name_in = args.docx
        worddoc.docx_name_out = args.word_out
        worddoc.convert()
    if args.lbnldevice is not None:
        worddoc = CreateWordDoc(lbnldevice=args.lbnldevice)
        worddoc.docx_name_in = args.docx
        worddoc.docx_name_out = args.word_out
        worddoc.convert()
        
except:
    #print ("error in ", args.json)
    traceback.print_exc()
    pass
    