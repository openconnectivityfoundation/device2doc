#############################
#
#    copyright 2016 Open Interconnect Consortium, Inc. All rights reserved.
#    Redistribution and use in source and binary forms, with or without modification,
#    are permitted provided that the following conditions are met:
#    1.  Redistributions of source code must retain the above copyright notice,
#        this list of conditions and the following disclaimer.
#    2.  Redistributions in binary form must reproduce the above copyright notice,
#        this list of conditions and the following disclaimer in the documentation and/or other materials provided
#        with the distribution.
#
#    THIS SOFTWARE IS PROVIDED BY THE OPEN INTERCONNECT CONSORTIUM, INC. "AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES,
#    INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE OR
#    WARRANTIES OF NON-INFRINGEMENT, ARE DISCLAIMED. IN NO EVENT SHALL THE OPEN INTERCONNECT CONSORTIUM, INC. OR
#    CONTRIBUTORS BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
#    (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS;
#    OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
#    OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE,
#    EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
#############################


import time
import os
import json
import random
import sys
import argparse
import traceback
from datetime import datetime
from time import gmtime, strftime
import jsonref

if sys.version_info < (3, 5):
    raise Exception("ERROR: Python 3.5 or more is required, you are currently running Python %d.%d!" %
                    (sys.version_info[0], sys.version_info[1]))
#
# docx imports
#
try:
    from docx import Document
except:
    print("missing docx:")
    print ("Trying to Install required module: python-docx (docx)")
    os.system('python -m pip install python-docx')
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def load_json_schema(filename, my_dir):
    """
    load the JSON schema file
    :param filename: filename (with extension)
    :param my_dir: path to the file
    :return: json_dict
    """
    full_path = os.path.join(my_dir, filename)
    if os.path.isfile(full_path) is False:
        print ("json file does not exist:", full_path)

    linestring = open(full_path, 'r').read()
    json_dict = json.loads(linestring)

    return json_dict


def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

def Table_annex(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table-Annex \* ARABIC  \s 9 '
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


class CreateWordDoc(object):
    def __init__(self, enum=None, docx_name_in=None, docx_name_out=None):
        """
        initialize the class


        """
        # input arguments
        self.docx_name_in = docx_name_in
        self.docx_name_out = docx_name_out

        # initialise the variable
        self.enum = enum
        self.enum_filename = enum

        schema_string = open(self.enum_filename, 'r', encoding='UTF8').read()
        json_dict = json.loads(schema_string)
        self.json_parse_tree = json_dict

    def swag_sanitize_description(self, description):
        """
        removes line breaks, quotes  etc
        :param description: input string
        :return: text string
        """
        text = description.replace("\n", "@cr").replace("'", "<COMMA>").replace('"', "<COMMA>")
        return text

    def swag_unsanitize_description(self, description):
        """
        removes line breaks, quotes  etc
        :param description: input string
        :return: text string
        """
        text = description.replace("@cr", "\n").replace("<COMMA>", "'")
        return text



    def enumerationdescriptions(self, parse_tree):
        """
        list all enumerations as a dashed list.
        :param parse_tree: json parse_tree of the enumeration set
        """
        # Add L2 Heading and lead in paragraphs
        heading = self.document.add_heading('Alphabetical list of standardized enumeration types', level=2)
        if self.annex_switch is True:
            heading.style = 'ANNEX-heading2'

        intro_para = self.document.add_paragraph('<Table Reference Here> lists the standardized enumeration types that may be present within Resource Properties where the Property is defined as containing values from this clause. The enumerations also apply to Semantic Tags where the tag is defined as containing values from this clause.')
        intro_para.style = 'PARAGRAPH'

        # create the caption
        paragraph = self.document.add_paragraph('Table ', style='Caption')
        if self.annex_switch is True:
            Table_annex (paragraph)
        else:
            Table (paragraph)

        paragraph.add_run(' – The defined set of standardized enumerations')
        paragraph.style = 'TABLE-title'

        self.tableAttribute = self.document.add_table(rows=1, cols=2, style='TABLE-A')
        self.tableAttribute.autofit = True
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = 'Enumeration'
        hdr_cells[0].paragraphs[0].style = 'TABLE-col-heading'
        hdr_cells[1].text = 'Description'
        hdr_cells[1].paragraphs[0].style = 'TABLE-col-heading'

        enumerationlist = parse_tree["supportedenumerations"]

        # Add the enumerations plus descriptions to the created Table
        for enumerationobject in enumerationlist:
            for enumname,enumdesc in enumerationobject.items():
                row_cells = self.tableAttribute.add_row().cells
                row_cells[0].text = enumname
                row_cells[0].paragraphs[0].style = 'TABLE-cell'
                row_cells[1].text = enumdesc
                row_cells[1].paragraphs[0].style = 'TABLE-cell'

        # Set the width of the first column
        for cell in self.tableAttribute.columns[0].cells:
            cell.width = 30

    def generate_sections(self, parse_tree):
        """
        generate the individual sections
        :param parse_tree:
        """
        self.list_devices(parse_tree)

    def convert(self):
        """
        conversion of the swagger data into the word document

        :return:
        """
        try:
            self.document = Document(docx=self.docx_name_in)
        except:
            print ("could not load file: ", self.docx_name_in)
            print ("make sure that docx file exist..")
            return

        if self.enum is not None:
            self.enumerationdescriptions(self.json_parse_tree)
        ### add here more conversions going forward..

        if self.docx_name_out is not None:
            self.document.save(self.docx_name_out)
            print ("document saved..", self.docx_name_out)


#
#   main of script
#
print ("************************")
print ("*** device2doc (v1) ***")
print ("************************")
parser = argparse.ArgumentParser()

parser.add_argument( "-ver", "--verbose", help="Execute in verbose mode", action='store_true')

parser.add_argument( "-enum", "--enum", default=None, help="enumeration file name (json)",  nargs='?', const="", required=False)
parser.add_argument( "-docx", "--docx", default=None, help="word file in",  nargs='?', const="", required=False)
parser.add_argument( "-word_out", "--word_out", default=None, help="word file out",  nargs='?', const="", required=False)
parser.add_argument("-annex", "--annex", default=None, help="uses a annex heading instead of normal heading (--annex true)")


args = parser.parse_args()


print("enum file     : " + str(args.enum))
print("docx            : " + str(args.docx))
print("word_out        : " + str(args.word_out))
print("annex       : " + str(args.annex))

print("")

try:
    if args.enum is not None:
        worddoc = CreateWordDoc(enum=args.enum)
        worddoc.docx_name_in = args.docx
        worddoc.docx_name_out = args.word_out

        annex_switch = args.annex
        if annex_switch is None:
            annex_switch = False
        else:
            annex_switch = True

        worddoc.annex_switch = annex_switch

        worddoc.convert()

except:
    #print ("error in ", args.json)
    traceback.print_exc()
    pass
